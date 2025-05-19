from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from . import logger_config  # Относительный импорт

logger = logger_config.setup_logger(__name__)

# --- Константы для стилей по умолчанию (если не найдены в word_styles.json) ---
DEFAULT_FONT = "Calibri"
DEFAULT_FONT_SIZE_PT_VAL = 11
TITLE_FONT_SIZE_PT_VAL = 18
H1_FONT_SIZE_PT_VAL = 14  # Для Заголовка 1
H2_FONT_SIZE_PT_VAL = 13  # Для Заголовка 2
H3_FONT_SIZE_PT_VAL = 12  # Для Заголовка 3
TASK_ITEM_FONT_SIZE_PT_VAL = 10  # Для текста задач в списках


# --- Вспомогательные функции ---

def _get_internal_name_from_mapping(fields_mapping_config: list, standard_csv_header: str, default_internal: str,
                                    alt_internals: list = None) -> str:
    """Ищет internal_name в fields_mapping_config."""
    if alt_internals is None: alt_internals = []
    for field_spec in fields_mapping_config:
        spec_internal = field_spec.get('internal_name')
        if spec_internal and (spec_internal == default_internal or spec_internal in alt_internals):
            return spec_internal
        spec_csv_header = field_spec.get('csv_header', '').lower()
        if spec_csv_header == standard_csv_header.lower():
            return field_spec.get('internal_name', field_spec.get('csv_header'))
    return default_internal


def set_run_font(run, font_name=None, size_pt_val=None, bold=None, italic=None, color_rgb=None):
    """Устанавливает свойства шрифта для объекта Run."""
    if font_name:
        run.font.name = font_name
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        r.rPr.rFonts.set(qn('w:cs'), font_name)
    if size_pt_val is not None:
        try:
            actual_size = Pt(int(size_pt_val))
            run.font.size = actual_size
            logger.debug(
                f"  [set_run_font] Установлен размер: {actual_size} для '{run.text[:20]}...' (исходное: {size_pt_val})")
        except ValueError:
            logger.warning(f"Неверное значение для размера шрифта: {size_pt_val}. Не удалось установить размер.")
        except Exception as e:
            logger.warning(f"Ошибка при установке размера шрифта ({size_pt_val}): {e}")
    if bold is not None: run.font.bold = bold
    if italic is not None: run.font.italic = italic
    if color_rgb is not None and isinstance(color_rgb, RGBColor): run.font.color.rgb = color_rgb


def add_styled_paragraph(document, text="", style_name=None, font_name=None, size_pt_val=None,
                         bold=None, italic=None, align=None,
                         space_after_pt=None, space_before_pt=None):
    """Добавляет обычный параграф с указанными стилями."""
    p = document.add_paragraph()
    if text:
        run = p.add_run(text)
        if font_name or size_pt_val is not None or bold is not None or italic is not None:
            set_run_font(run, font_name, size_pt_val, bold, italic)

    if style_name:  # Применяем стиль параграфа Word, если указан
        try:
            p.style = style_name
        except Exception as e:
            logger.warning(f"Не удалось применить стиль параграфа '{style_name}': {e}")

    pf = p.paragraph_format
    if align: pf.alignment = align
    if space_after_pt is not None: pf.space_after = Pt(int(space_after_pt))
    if space_before_pt is not None: pf.space_before = Pt(int(space_before_pt))
    return p


def _apply_field_style_to_run(run, field_style: dict, default_font_name: str, default_font_size: int):
    """Применяет стили из field_style к объекту run."""
    set_run_font(
        run,
        font_name=field_style.get('font_name', default_font_name),
        size_pt_val=field_style.get('font_size', default_font_size),
        bold=field_style.get('bold', False),
        italic=field_style.get('italic', False)
    )


def _add_task_fields_to_paragraph(paragraph, task_dict: dict, fields_to_display: list,
                                  default_font_name: str, default_font_size: int,
                                  style_key_in_field_spec: str):
    """
    Добавляет отформатированные поля задачи в существующий параграф.
    Версия "v1.0" до сложных исправлений многострочности.
    """
    content_added_to_paragraph_this_field_iteration = False

    for field_spec in fields_to_display:
        internal_name = field_spec.get('internal_name')
        value = task_dict.get(internal_name)
        original_value_str = str(value) if value is not None else ""
        stripped_value_str = original_value_str.strip()

        if not stripped_value_str:
            if internal_name == 'task_report_text':
                original_value_str = "Нет описания."
                stripped_value_str = original_value_str
            elif internal_name == 'setup_instructions':
                original_value_str = "Инструкции отсутствуют."
                stripped_value_str = original_value_str
            else:
                logger.debug(f"Поле '{internal_name}' пустое или состоит из пробелов, пропускаем.")
                continue

        style_options = field_spec.get(style_key_in_field_spec, {})

        if content_added_to_paragraph_this_field_iteration and style_options.get('new_line_before', False):
            paragraph.add_run().add_break()

        prefix = style_options.get('prefix', '')
        if prefix:
            _apply_field_style_to_run(paragraph.add_run(prefix), style_options, default_font_name, default_font_size)
            content_added_to_paragraph_this_field_iteration = True

        label = field_spec.get('report_label', '')
        if label:
            _apply_field_style_to_run(paragraph.add_run(label + (" " if stripped_value_str else "")), style_options,
                                      default_font_name, default_font_size)
            content_added_to_paragraph_this_field_iteration = True

        if style_options.get('multiline', False) and '\n' in original_value_str:
            lines = [line for line in original_value_str.splitlines() if line.strip()]
            for i, line_text in enumerate(lines):
                if i > 0: paragraph.add_run().add_break()
                _apply_field_style_to_run(paragraph.add_run(line_text), style_options, default_font_name,
                                          default_font_size)
                content_added_to_paragraph_this_field_iteration = True
        elif stripped_value_str:
            _apply_field_style_to_run(paragraph.add_run(stripped_value_str), style_options, default_font_name,
                                      default_font_size)
            content_added_to_paragraph_this_field_iteration = True

        suffix = style_options.get('suffix', '')
        if suffix:
            _apply_field_style_to_run(paragraph.add_run(suffix), style_options, default_font_name, default_font_size)
            content_added_to_paragraph_this_field_iteration = True


# --- Основные функции генерации секций ---

def create_title_section(document, report_title, logo_path, styles_config):
    logger.info("Создание титульной секции...")
    fonts = styles_config.get('fonts', {})
    font_sizes = styles_config.get('font_sizes', {})
    para_spacing = styles_config.get('paragraph_spacing', {})

    if logo_path and os.path.exists(logo_path):
        try:
            document.add_picture(logo_path, width=Inches(1.5))
            logger.info(f"Логотип добавлен из: {logo_path}")
        except Exception as e:
            logger.error(f"Не удалось добавить логотип из {logo_path}: {e}", exc_info=True)
    elif logo_path:
        logger.warning(f"Файл логотипа не найден: {logo_path}")

    title_paragraph = document.add_heading(report_title, level=0)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if title_paragraph.runs:
        title_font_name_cfg = fonts.get('title', None)
        title_font_size_cfg = font_sizes.get('title', None)
        set_run_font(title_paragraph.runs[0],
                     font_name=title_font_name_cfg,
                     size_pt_val=title_font_size_cfg,
                     bold=True)

    pf = title_paragraph.paragraph_format
    pf.space_after = Pt(int(para_spacing.get('after_title', 24)))
    pf.space_before = Pt(int(para_spacing.get('before_title', 12)))

    logger.info(f"Заголовок отчета добавлен (Уровень 0): '{report_title}'")


def create_microservices_version_table(document, versions_data, styles_config):
    if not versions_data:
        logger.warning("Нет данных для таблицы версий МС. Таблица не будет создана.")
        return
    logger.info("Создание таблицы версий МС...")
    fonts = styles_config.get('fonts', {});
    font_sizes = styles_config.get('font_sizes', {})
    colors = styles_config.get('colors_hex', {});
    table_props = styles_config.get('table_properties', {})
    default_font = fonts.get('default', DEFAULT_FONT)
    table_header_font_size_val = font_sizes.get('table_header', 10)
    table_content_font_size_val = font_sizes.get('table_content', 10)
    table_header_bg_color_hex = colors.get('table_header_background', "D9D9D9")

    heading_text_table = "Версии компонентов релиза:"
    h2_table_title = document.add_heading(heading_text_table, level=2)
    if h2_table_title.runs:
        set_run_font(h2_table_title.runs[0],
                     font_name=fonts.get('heading2', None),
                     size_pt_val=font_sizes.get('heading2', None),
                     bold=True)
    pf_h2_table = h2_table_title.paragraph_format
    pf_h2_table.space_before = Pt(int(styles_config.get('paragraph_spacing', {}).get('before_heading2', 12)))
    pf_h2_table.space_after = Pt(int(styles_config.get('paragraph_spacing', {}).get('after_heading2_table', 6)))

    try:
        table = document.add_table(rows=1, cols=2);
        table.style = 'TableGrid'
    except Exception as e:
        logger.error(f"Не удалось создать таблицу: {e}", exc_info=True); return

    header_texts = ['Микросервис', 'Версия'];
    hdr_cells = table.rows[0].cells
    for i, cell_obj in enumerate(hdr_cells):
        try:
            cell_obj.text = header_texts[i];
            p = cell_obj.paragraphs[0];
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for old_run in p.runs: p._element.remove(old_run._element)
            run = p.add_run(header_texts[i]);
            set_run_font(run, font_name=default_font, size_pt_val=table_header_font_size_val, bold=True)
            if table_header_bg_color_hex:
                shd = OxmlElement('w:shd');
                shd.set(qn('w:val'), 'clear');
                shd.set(qn('w:color'), 'auto');
                shd.set(qn('w:fill'), table_header_bg_color_hex)
                cell_obj._tc.get_or_add_tcPr().append(shd)
        except Exception as e:
            logger.warning(f"Ошибка стилизации заголовка таблицы '{header_texts[i]}': {e}", exc_info=True)

    for item in versions_data:
        try:
            row_cells = table.add_row().cells;
            cell_texts = [item.get('microservice', ''), item.get('version', '')]
            for i, cell_obj in enumerate(row_cells):
                cell_obj.text = cell_texts[i];
                p = cell_obj.paragraphs[0]
                for old_run in p.runs: p._element.remove(old_run._element)
                run = p.add_run(cell_texts[i]);
                set_run_font(run, font_name=default_font, size_pt_val=table_content_font_size_val)
        except Exception as e:
            logger.warning(f"Ошибка доб. строки в таблицу ({item}): {e}", exc_info=True)

    try:
        col1_w_p = table_props.get('width_col1_percent', 30);
        col2_w_p = table_props.get('width_col2_percent', 70)
        assumed_total_width_inches = 6.0
        logger.debug(f"Исп. предпол. общая ширина для таблицы: {assumed_total_width_inches:.2f} дюймов.")
        c0w = assumed_total_width_inches * (col1_w_p / 100.0);
        c1w = assumed_total_width_inches * (col2_w_p / 100.0)
        if table.columns and len(table.columns) == 2:
            if c0w > 0.1:
                table.columns[0].width = Inches(c0w)
            else:
                logger.warning(f"Расч. ширина кол.0 ({c0w:.2f}) слишком мала.")
            if c1w > 0.1:
                table.columns[1].width = Inches(c1w)
            else:
                logger.warning(f"Расч. ширина кол.1 ({c1w:.2f}) слишком мала.")
        else:
            logger.warning("Нет доступа к колонкам таблицы для уст. ширины.")
    except Exception as e:
        logger.error(f"Крит. ошибка при уст. ширины колонок: {e}", exc_info=True)

    document.add_paragraph()
    logger.info("Таблица версий МС создана.")


def create_changes_section(document, grouped_data, styles_config, fields_mapping_config, report_config):
    logger.info("Создание раздела 'Перечень изменений'...")
    fonts = styles_config.get('fonts', {});
    font_sizes = styles_config.get('font_sizes', {})
    para_spacing = styles_config.get('paragraph_spacing', {});
    section_titles = report_config.get('report_section_titles', {})
    default_font = fonts.get('default', DEFAULT_FONT)
    task_font_size = font_sizes.get('task_item', TASK_ITEM_FONT_SIZE_PT_VAL)

    if not grouped_data:
        logger.warning("Нет данных для 'Перечня изменений'.")
        no_changes_text = report_config.get('report_section_titles', {}).get('no_changes_text',
                                                                             "Изменений в данной версии не зарегистрировано.")
        p_no_changes = document.add_paragraph()
        set_run_font(p_no_changes.add_run(no_changes_text), font_name=default_font, size_pt_val=task_font_size,
                     italic=True)
        pf = p_no_changes.paragraph_format
        pf.space_before = Pt(int(para_spacing.get('list_item_before', 6)))
        pf.space_after = Pt(int(para_spacing.get('list_item_after', 6)))
        return

    changes_title_text = section_titles.get('main_changes', "Перечень изменений")
    h1_changes = document.add_heading(changes_title_text, level=1)
    if h1_changes.runs:
        set_run_font(h1_changes.runs[0], font_name=fonts.get('heading1', None),
                     size_pt_val=font_sizes.get('heading1', None), bold=True)
    pf_h1c = h1_changes.paragraph_format
    pf_h1c.space_before = Pt(int(para_spacing.get('before_heading1', para_spacing.get('after_title', 24))))
    pf_h1c.space_after = Pt(int(para_spacing.get('after_heading1', 12)))

    fields_for_display = sorted(
        [f for f in fields_mapping_config if f.get('display_in_changes', False)],
        key=lambda f: f.get('changes_order', 99)
    )
    if not fields_for_display:
        logger.warning("Не настроены поля для отображения в 'Перечне изменений'.")

    key_internal_name_for_fallback = _get_internal_name_from_mapping(fields_mapping_config, 'issue key', 'issue_key')

    for ms_name, types_dict in grouped_data.items():
        h2_ms = document.add_heading(ms_name, level=2)
        if h2_ms.runs:
            set_run_font(h2_ms.runs[0], font_name=fonts.get('heading2', None),
                         size_pt_val=font_sizes.get('heading2', None), bold=True)
        pf_h2ms = h2_ms.paragraph_format
        pf_h2ms.space_before = Pt(int(para_spacing.get('before_heading2', 8)))
        pf_h2ms.space_after = Pt(int(para_spacing.get('after_heading2', 4)))

        for issue_type, tasks_list in types_dict.items():
            h3_type = document.add_heading(issue_type, level=3)
            if h3_type.runs:
                set_run_font(h3_type.runs[0], font_name=fonts.get('heading3', None),
                             size_pt_val=font_sizes.get('heading3', None), italic=True)
            pf_h3t = h3_type.paragraph_format
            pf_h3t.space_before = Pt(int(para_spacing.get('before_heading3', 4)))
            pf_h3t.space_after = Pt(int(para_spacing.get('after_heading3', 2)))

            if not tasks_list:
                try:
                    p_no = document.add_paragraph(style='ListBullet')
                except KeyError:
                    p_no = document.add_paragraph()
                set_run_font(p_no.add_run("Нет задач этого типа."), font_name=default_font, size_pt_val=task_font_size,
                             italic=True)
                pf = p_no.paragraph_format;
                pf.space_before = Pt(0);
                pf.space_after = Pt(int(para_spacing.get('list_item_after', 6)))
                continue

            for task_dict in tasks_list:
                current_list_style = 'ListBullet'
                try:
                    p_task = document.add_paragraph(style=current_list_style)
                except KeyError:
                    logger.warning(f"Стиль '{current_list_style}' не найден.");
                    p_task = document.add_paragraph()

                if fields_for_display:
                    _add_task_fields_to_paragraph(p_task, task_dict, fields_for_display,
                                                  default_font, task_font_size, 'changes_style')
                else:
                    key_val = task_dict.get(key_internal_name_for_fallback, "")
                    desc_val = task_dict.get('task_report_text', "Нет текста.")
                    if key_val: set_run_font(p_task.add_run(f"{key_val}: "), font_name=default_font,
                                             size_pt_val=task_font_size, bold=True)
                    set_run_font(p_task.add_run(desc_val), font_name=default_font, size_pt_val=task_font_size)

                pf = p_task.paragraph_format
                pf.space_before = Pt(int(para_spacing.get('list_item_before', 0)))
                pf.space_after = Pt(int(para_spacing.get('list_item_after', 6)))
    logger.info("Раздел 'Перечень изменений' создан.")


def create_setup_section(document, grouped_setup_data, styles_config, fields_mapping_config, report_config):
    logger.info("--- НАЧАЛО: Создание раздела 'Настройки системы' ---")
    logger.debug(f"Полученные grouped_setup_data (тип: {type(grouped_setup_data)}): "
                 f"{grouped_setup_data if grouped_setup_data is not None and len(grouped_setup_data) < 5 else 'Данные присутствуют или слишком большие для лога'}")
    if grouped_setup_data is not None and isinstance(grouped_setup_data, dict):
        logger.debug(f"Ключи в grouped_setup_data: {list(grouped_setup_data.keys())}")

    fonts = styles_config.get('fonts', {});
    font_sizes = styles_config.get('font_sizes', {})
    para_spacing = styles_config.get('paragraph_spacing', {});
    section_titles = report_config.get('report_section_titles', {})
    default_font = fonts.get('default', DEFAULT_FONT)
    task_font_size = font_sizes.get('task_item', TASK_ITEM_FONT_SIZE_PT_VAL)

    if not grouped_setup_data:
        logger.info("Нет данных для 'Настроек системы' (grouped_setup_data пуст или None). Раздел будет пропущен.")
        return

    setup_title_text = section_titles.get('system_setup', "Настройки системы (Заголовок по умолчанию)")
    logger.debug(f"Заголовок для раздела Настройки: '{setup_title_text}'")

    if not setup_title_text or not setup_title_text.strip():
        logger.warning("Текст заголовка для раздела 'Настройки системы' пуст. Заголовок H1 не будет добавлен.")
    else:
        h1_setup = document.add_heading(setup_title_text, level=1)
        logger.info(f"Добавлен заголовок H1 для Настроек: '{setup_title_text}'")
        if h1_setup.runs:
            h1_font_cfg_setup = fonts.get('heading1', None)
            h1_size_cfg_setup = font_sizes.get('heading1', None)
            logger.debug(f"Стилизация H1 для Настроек: шрифт={h1_font_cfg_setup}, размер={h1_size_cfg_setup}")
            set_run_font(h1_setup.runs[0], font_name=h1_font_cfg_setup, size_pt_val=h1_size_cfg_setup, bold=True)
        else:
            logger.warning(f"У заголовка H1 (Настройки) '{setup_title_text}' нет runs для стилизации.")

        pf_h1s = h1_setup.paragraph_format
        pf_h1s.space_before = Pt(int(para_spacing.get('before_heading1', para_spacing.get('after_title', 24))))
        pf_h1s.space_after = Pt(int(para_spacing.get('after_heading1', 12)))

    fields_for_setup_display = sorted(
        [f for f in fields_mapping_config if f.get('display_in_setup', False)],
        key=lambda f: f.get('setup_order', 99)
    )
    if not fields_for_setup_display:
        logger.warning(
            "Не настроены поля для отображения в 'Настройках системы' (display_in_setup: true). Задачи могут быть не отображены или отображены некорректно.")

    key_internal_name_fb_setup = _get_internal_name_from_mapping(fields_mapping_config, 'issue key', 'issue_key')
    summary_internal_name_fb_setup = _get_internal_name_from_mapping(fields_mapping_config, 'summary', 'summary_text',
                                                                     ['summary'])
    setup_instr_internal_name_fb_setup = _get_internal_name_from_mapping(fields_mapping_config,
                                                                         'custom field (инструкция по установке)',
                                                                         'setup_instructions')

    logger.debug(f"Количество микросервисов для настроек: {len(grouped_setup_data) if grouped_setup_data else 0}")
    for ms_name, tasks_list in grouped_setup_data.items():
        logger.debug(f"Обработка микросервиса для настроек: '{ms_name}'")
        if not ms_name or not ms_name.strip():
            logger.warning("Имя микросервиса для настроек пустое. Заголовок H2 не будет добавлен.")
        else:
            h2_ms_setup = document.add_heading(ms_name, level=2)
            logger.info(f"Добавлен заголовок H2 для Настроек (МС): '{ms_name}'")
            if h2_ms_setup.runs:
                h2_font_cfg_setup = fonts.get('heading2', None)
                h2_size_cfg_setup = font_sizes.get('heading2', None)
                logger.debug(
                    f"Стилизация H2 для Настроек (МС '{ms_name}'): шрифт={h2_font_cfg_setup}, размер={h2_size_cfg_setup}")
                set_run_font(h2_ms_setup.runs[0], font_name=h2_font_cfg_setup, size_pt_val=h2_size_cfg_setup, bold=True)
            else:
                logger.warning(f"У заголовка H2 (Настройки, МС '{ms_name}') нет runs для стилизации.")

            pf_h2ms_s = h2_ms_setup.paragraph_format
            pf_h2ms_s.space_before = Pt(int(para_spacing.get('before_heading2', 8)))
            pf_h2ms_s.space_after = Pt(int(para_spacing.get('after_heading2', 4)))

        if not tasks_list:
            logger.debug(f"Для МС '{ms_name}' (Настройки) нет задач с инструкциями.")
            try:
                p_no = document.add_paragraph(style='ListBullet')
            except KeyError:
                p_no = document.add_paragraph()
            set_run_font(p_no.add_run("Нет инструкций по настройке для этого компонента."),
                         font_name=default_font, size_pt_val=task_font_size, italic=True)
            pf = p_no.paragraph_format
            pf.space_before = Pt(int(para_spacing.get('list_item_before', 0)))
            pf.space_after = Pt(int(para_spacing.get('list_item_after', 6)))
            continue

        logger.debug(f"Для МС '{ms_name}' (Настройки) найдено задач: {len(tasks_list)}")
        for task_dict in tasks_list:
            current_list_style = 'ListBullet'
            try:
                p_task = document.add_paragraph(style=current_list_style)
                logger.debug(
                    f"  Добавлен параграф для задачи (Настройки, стиль {current_list_style}): {task_dict.get(key_internal_name_fb_setup, 'ID?')} ")
            except KeyError:
                logger.warning(f"Стиль '{current_list_style}' не найден. Используется обычный параграф.")
                p_task = document.add_paragraph()

            if fields_for_setup_display:
                _add_task_fields_to_paragraph(p_task, task_dict, fields_for_setup_display,
                                              default_font, task_font_size, 'setup_style')
            else:
                logger.debug(
                    f"  Используется фоллбэк для полей задачи {task_dict.get(key_internal_name_fb_setup, 'ID?')} в 'Настройках системы'.")
                key_val = task_dict.get(key_internal_name_fb_setup, "")
                summary_val = task_dict.get(summary_internal_name_fb_setup, "")
                instr_val = task_dict.get(setup_instr_internal_name_fb_setup, "Инструкции отсутствуют.")

                header_parts_fb = [p for p in [key_val, summary_val] if p]
                header_text_fb = ": ".join(header_parts_fb) if header_parts_fb else "Инструкция"

                run_h = p_task.add_run(header_text_fb)
                set_run_font(run_h, font_name=default_font, size_pt_val=task_font_size, bold=True)

                if instr_val and instr_val != "Инструкции отсутствуют.":
                    if header_text_fb and header_text_fb != "Инструкция":
                        p_task.add_run().add_break()

                    instr_lines_fb = instr_val.splitlines()
                    for i_fb, line_fb in enumerate(instr_lines_fb):
                        if i_fb > 0 and line_fb.strip(): p_task.add_run().add_break()
                        if line_fb.strip():
                            set_run_font(p_task.add_run(line_fb), font_name=default_font, size_pt_val=task_font_size)

            pf = p_task.paragraph_format
            pf.space_before = Pt(int(para_spacing.get('list_item_before', 0)))
            pf.space_after = Pt(int(para_spacing.get('list_item_after', 6)))
    logger.info("--- ЗАВЕРШЕНИЕ: Создание раздела 'Настройки системы' ---")


def generate_report_docx(output_filename, report_title_text, logo_full_path,
                         microservice_versions_list, word_styles_config,
                         grouped_data_for_changes,
                         grouped_data_for_setup,  # Это правильное имя параметра
                         main_config_for_titles,
                         fields_mapping_for_details
                         ):
    logger.info(f"Начало генерации DOCX отчета: {output_filename}")
    doc = Document()

    create_title_section(doc, report_title_text, logo_full_path, word_styles_config)
    create_microservices_version_table(doc, microservice_versions_list, word_styles_config)
    create_changes_section(doc, grouped_data_for_changes, word_styles_config, fields_mapping_for_details,
                           main_config_for_titles)

    logger.info("--- ПЕРЕД ВЫЗОВОМ create_setup_section В generate_report_docx ---")
    # ИСПРАВЛЕНИЕ ЗДЕСЬ: используем grouped_data_for_setup (имя параметра функции)
    logger.debug(f"Тип grouped_data_for_setup: {type(grouped_data_for_setup)}")
    if grouped_data_for_setup is not None:
        keys_to_log = list(grouped_data_for_setup.keys())[:2] if isinstance(grouped_data_for_setup,
                                                                            dict) else "Не словарь или пуст"
        logger.debug(
            f"Содержимое grouped_data_for_setup (первые ключи, если есть): {keys_to_log}, Количество элементов: {len(grouped_data_for_setup) if hasattr(grouped_data_for_setup, '__len__') else 'N/A'}")
    else:
        logger.debug("grouped_data_for_setup is None (передан в generate_report_docx)")

    create_setup_section(doc, grouped_data_for_setup, word_styles_config, fields_mapping_for_details,
                         main_config_for_titles)
    logger.info("--- ПОСЛЕ ВЫЗОВА create_setup_section В generate_report_docx ---")

    try:
        output_dir = os.path.dirname(output_filename)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir);
            logger.info(f"Создана директория: {output_dir}")
        doc.save(output_filename)
        logger.info(f"Отчет успешно сохранен: {output_filename}")
    except Exception as e:
        logger.error(f"Не удалось сохранить документ {output_filename}: {e}", exc_info=True)